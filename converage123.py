# -*- coding: utf-8 -*-
import re, csv, json, sys
from pathlib import Path

# 字符范围
KANJI    = r"[一-龯々〆ヶ]"
KANA     = r"[ぁ-ゖァ-ヺー･・]"   # 平/片假名 + 长音符等
KANA_WS = rf"(?:{KANA}|[\s　])+"  # 假名或空白（含全角空格）

# —— 基础工具 ——
def katakana_to_hiragana(s: str) -> str:
    out = []
    for ch in s:
        o = ord(ch)
        if 0x30A1 <= o <= 0x30FA:  # カタカナ → ひらがな
            out.append(chr(o - 0x60))
        else:
            out.append(ch)
    return ''.join(out)

# —— 剥离“说话人标注”（整段裁掉），示例：鯉夏（こいなつ）：…  / 鯉夏(こいなつ): …
SPEAKER_HEAD = re.compile(
        rf"^[（][\S, ]+[）]", 
        flags=re.MULTILINE
)

# —— 剥离“假名注音”只留汉字，支持全/半角括号：漢字（かな）/漢字(かな) → 漢字
INLINE_FURI_FULL = re.compile(rf"({KANJI}+)[（]({KANA_WS})[）]")
INLINE_FURI_HALF = re.compile(rf"({KANJI}+)\(({KANA_WS})\)")

def strip_speaker_and_furigana_text(text: str) -> str:
    # 1) 行首说话人标注整段移除（只要括号里是纯假名就删）
    text = SPEAKER_HEAD.sub("", text)
    # 2) 正文内 furigana 移除，只留汉字
    text = INLINE_FURI_FULL.sub(r"\1", text)
    text = INLINE_FURI_HALF.sub(r"\1", text)
    return text

# —— SRT 剥离：保留时间轴/行号，仅改台词行 ——
def strip_srt(path_in: Path, path_out: Path):
    raw = path_in.read_text(encoding="utf-8", errors="ignore")
    blocks = re.split(r"\r?\n\r?\n", raw.strip())
    new_blocks = []
    for b in blocks:
        lines = b.splitlines()
        if not lines: 
            continue
        # 找时间行
        time_idx = None
        for i, ln in enumerate(lines[:4]):
            if "-->" in ln:
                time_idx = i; break
        if time_idx is None:
            # 不是标准 SRT 块：整块照剥离规则处理
            new_blocks.append(strip_speaker_and_furigana_text(b))
            continue
        head = lines[:time_idx+1]
        body = lines[time_idx+1:]
        new_body = [strip_speaker_and_furigana_text(ln) for ln in body]
        new_blocks.append("\n".join(head + new_body))
    path_out.write_text("\n\n".join(new_blocks) + "\n", encoding="utf-8")

# —— ASS 剥离：只改“Text”字段（第10段），其他字段不动 ——
def strip_ass(path_in: Path, path_out: Path):
    out_lines = []
    for line in path_in.read_text(encoding="utf-8", errors="ignore").splitlines():
        if not line.startswith("Dialogue:"):
            out_lines.append(line); continue
        # ASS: Dialogue: Layer,Start,End,Style,Name,MarginL,MarginR,MarginV,Effect,Text
        parts = line.split(",", 9)
        if len(parts) < 10:
            out_lines.append(line); continue
        head = ",".join(parts[:9])
        text = parts[9]
        # 去除 ASS 覆盖标签再剥离，最后保留原标签次序（这里简单做：仅剥离文本，不改标签）
        # 若需移除形如 {\i1} 的标签，可先用 re.sub(r"\{\\[^}]+\}", "", text)
        cleaned_text = strip_speaker_and_furigana_text(text)
        out_lines.append(head + "," + cleaned_text)
    path_out.write_text("\n".join(out_lines) + "\n", encoding="utf-8")

def simple_tokenize(s: str):
    """使用 jaconv 和分词来计算 token 数量"""
    s = s.strip()
    if not s: return []
    
    # 如果有空格，按空格分词
    if re.search(r"\s", s):
        return re.split(r"\s+", s)
    
    # 日文分词处理
    try:
        import jaconv
        import fugashi
        
        # 片假名转平假名
        text = jaconv.kata2hira(s)
        
        # 使用 fugashi 分词
        tagger = fugashi.Tagger()
        tokens = []
        for word in tagger(text):
            if word.surface.strip():  # 只保留非空的词
                tokens.append(word.surface)
        
        return tokens if tokens else list(s)  # 如果分词失败，回退到字符级别
        
    except ImportError:
        # 如果没有安装相关库，回退到字符级别
        return list(s)

def analyze_docx_yellow(docx_path: Path, csv_out: Path=None):
    """
    统计 Word(.docx) 黄色高亮覆盖率。
    复用已有的文本清洗函数：strip_speaker_and_furigana_text
    返回字符/简易token两套覆盖率，并可导出清洗后的高亮片段 CSV。
    """
    from docx import Document
    from docx.enum.text import WD_COLOR_INDEX
    import csv

    doc = Document(str(docx_path))

    total_chars = total_tokens = 0
    hilite_chars = hilite_tokens = 0
    rows = []

    for para in doc.paragraphs:
        for run in para.runs:
            raw = run.text or ""
            # ★ 复用已有的清洗函数
            txt = strip_speaker_and_furigana_text(raw)
            if not txt:
                continue

            tchars = len(txt)
            ttoks  = len(simple_tokenize(txt))
            total_chars  += tchars
            total_tokens += ttoks

            if run.font.highlight_color == WD_COLOR_INDEX.YELLOW:
                hilite_chars  += tchars
                hilite_tokens += ttoks
                # 只收非空串，避免写入空白行
                rows.append([txt])

    stats = {
        "total_chars": total_chars,
        "hilite_chars": hilite_chars,
        "char_coverage": (hilite_chars / total_chars) if total_chars else 0.0,
        "total_tokens": total_tokens,
        "hilite_tokens": hilite_tokens,
        "token_coverage": (hilite_tokens / total_tokens) if total_tokens else 0.0,
    }

    if csv_out:
        with open(csv_out, 'w', encoding='utf-8', newline='') as f:
            w = csv.writer(f)
            w.writerow(['highlight_text_cleaned'])
            w.writerows(rows)

    return stats

# —— 导出对白为纯文本 ——
def extract_dialogue_from_srt(srt_path: Path, txt_out: Path, clean: bool = True):
    """从SRT字幕文件提取对白，导出为纯文本"""
    raw = srt_path.read_text(encoding="utf-8", errors="ignore")
    blocks = re.split(r"\r?\n\r?\n", raw.strip())
    
    dialogues = []
    for b in blocks:
        lines = b.splitlines()
        if not lines: 
            continue
        # 找时间行
        time_idx = None
        for i, ln in enumerate(lines[:4]):
            if "-->" in ln:
                time_idx = i; break
        if time_idx is None:
            continue
        
        # 提取台词（时间行之后的所有行）
        dialogue_lines = lines[time_idx+1:]
        for line in dialogue_lines:
            if line.strip():
                if clean:
                    cleaned = strip_speaker_and_furigana_text(line.strip())
                    if cleaned:
                        dialogues.append(cleaned)
                else:
                    dialogues.append(line.strip())
    
    txt_out.write_text("\n".join(dialogues) + "\n", encoding="utf-8")
    return len(dialogues)

def extract_dialogue_from_ass(ass_path: Path, txt_out: Path, clean: bool = True):
    """从ASS字幕文件提取对白，导出为纯文本"""
    dialogues = []
    for line in ass_path.read_text(encoding="utf-8", errors="ignore").splitlines():
        if not line.startswith("Dialogue:"):
            continue
        # ASS: Dialogue: Layer,Start,End,Style,Name,MarginL,MarginR,MarginV,Effect,Text
        parts = line.split(",", 9)
        if len(parts) < 10:
            continue
        text = parts[9].strip()
        
        # 移除ASS标签
        text = re.sub(r"\{[^}]*\}", "", text)
        
        if text:
            if clean:
                cleaned = strip_speaker_and_furigana_text(text)
                if cleaned:
                    dialogues.append(cleaned)
            else:
                dialogues.append(text)
    
    txt_out.write_text("\n".join(dialogues) + "\n", encoding="utf-8")
    return len(dialogues)

def extract_dialogue_from_docx(docx_path: Path, txt_out: Path, clean: bool = True):
    """从Word文档提取所有文本，导出为纯文本"""
    from docx import Document
    
    doc = Document(str(docx_path))
    dialogues = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            if clean:
                cleaned = strip_speaker_and_furigana_text(text)
                if cleaned:
                    dialogues.append(cleaned)
            else:
                dialogues.append(text)
    
    txt_out.write_text("\n".join(dialogues) + "\n", encoding="utf-8")
    return len(dialogues)

# —— Word文档清理：直接修改文档内容，保留样式 ——
def strip_docx(docx_path: Path, docx_out: Path):
    """清理Word文档：移除说话人标注和假名注音，保存为新文档，保留原有样式"""
    from docx import Document
    import re
    
    doc = Document(str(docx_path))
    
    # Word专用的清理正则表达式
    speaker_patterns = [
        # 匹配行首的全角括号说话人标注（包括复杂和简单格式）
        re.compile(r'^[（][^）]*[）]', flags=re.MULTILINE),
        # 匹配行首的半角括号说话人标注
        re.compile(r'^\([^)]*\)', flags=re.MULTILINE),
    ]
    
    # 假名注音清理
    kanji_pattern = r"[一-龯々〆ヶ]"
    kana_pattern = r"[ぁ-ゖァ-ヺー･・]"
    kana_ws_pattern = rf"(?:{kana_pattern}|[\s　])+"
    
    furi_full = re.compile(rf"({kanji_pattern}+)[（]({kana_ws_pattern})[）]")
    furi_half = re.compile(rf"({kanji_pattern}+)\(({kana_ws_pattern})\)")
    
    def clean_run_text(run):
        """清理单个run的文本，保留样式"""
        if not run.text.strip():
            return
        
        text = run.text
        
        # 1. 移除说话人标注
        for pattern in speaker_patterns:
            text = pattern.sub("", text)
        
        # 2. 移除假名注音，只保留汉字
        text = furi_full.sub(r"\1", text)
        text = furi_half.sub(r"\1", text)
        
        # 3. 清理多余空白
        text = re.sub(r'\n\s*\n', '\n', text)  # 多个空行变一个
        text = re.sub(r'^\s+', '', text, flags=re.MULTILINE)  # 行首空白
        text = text.strip()
        
        # 更新run的文本
        run.text = text
    
    # 处理段落中的文本
    for para in doc.paragraphs:
        for run in para.runs:
            clean_run_text(run)
    
    # 处理表格中的文本
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        clean_run_text(run)
    
    # 处理页眉页脚
    for section in doc.sections:
        # 页眉
        if section.header:
            for para in section.header.paragraphs:
                for run in para.runs:
                    clean_run_text(run)
        # 页脚
        if section.footer:
            for para in section.footer.paragraphs:
                for run in para.runs:
                    clean_run_text(run)
    
    doc.save(str(docx_out))
    return True

# —— 测试清理效果 ——
def test_speaker_cleaning():
    """测试说话人清理效果"""
    test_cases = [
        "（炭治郎・嘴平伊之助(はしびら いのすけ)）ん？",
        "（善逸(ぜんいつ)）ええ… あっ ねえ？", 
        "（宇髄天元(うずい てんげん)）いいか！",
        "（善逸）简单格式测试",  # 添加简单格式测试
        "普通の文章(ふつう)です。",  # 这个不应该被清理
    ]
    
    import re
    
    # Word专用的清理正则表达式
    speaker_patterns = [
        # 匹配行首的全角括号说话人标注（包括复杂和简单格式）
        re.compile(r'^[（][^）]*[）]', flags=re.MULTILINE),
        # 匹配行首的半角括号说话人标注
        re.compile(r'^\([^)]*\)', flags=re.MULTILINE),
    ]
    
    # 假名注音清理
    kanji_pattern = r"[一-龯々〆ヶ]"
    kana_pattern = r"[ぁ-ゖァ-ヺー･・]"
    kana_ws_pattern = rf"(?:{kana_pattern}|[\s　])+"
    
    furi_full = re.compile(rf"({kanji_pattern}+)[（]({kana_ws_pattern})[）]")
    furi_half = re.compile(rf"({kanji_pattern}+)\(({kana_ws_pattern})\)")
    
    for test_text in test_cases:
        print(f"\n原文: {test_text}")
        
        result = test_text
        
        # 1. 移除说话人标注
        for pattern in speaker_patterns:
            result = pattern.sub("", result)
        
        # 2. 移除假名注音，只保留汉字
        result = furi_full.sub(r"\1", result)
        result = furi_half.sub(r"\1", result)
        
        # 3. 清理多余空白
        result = re.sub(r'\n\s*\n', '\n', result)
        result = re.sub(r'^\s+', '', result, flags=re.MULTILINE)
        result = result.strip()
        
        print(f"清理后: {result}")
    
    return True


def main():
    import argparse
    p = argparse.ArgumentParser()
    p.add_argument("--strip_srt", type=str, help="剥离 SRT：移除行首说话人 + 假名注音，仅改台词文本")
    p.add_argument("--strip_ass", type=str, help="剥离 ASS：移除台词文本中的说话人 + 假名注音")
    p.add_argument("--strip_docx", type=str, help="剥离 Word文档：移除说话人标注和假名注音")
    p.add_argument("--docx", type=str, help="统计 Word 黄色高亮覆盖率")
    p.add_argument("--csv",  type=str, help="导出高亮文本 CSV")
    
    # 统一的txt导出参数
    p.add_argument("--txt", action="store_true", help="同时导出对白为txt文件")
    p.add_argument("--no_clean", action="store_true", help="导出txt时不清理说话人和假名注音（保持原文）")
    p.add_argument("--txt_out", type=str, help="指定txt输出文件路径（可选）")
    
    # 测试参数
    p.add_argument("--test_clean", action="store_true", help="测试说话人清理效果")
    
    args = p.parse_args()

    # 测试清理效果
    if args.test_clean:
        test_speaker_cleaning()
        return

    clean_text = not args.no_clean

    if args.strip_srt:
        pin  = Path(args.strip_srt)
        pout = pin.with_suffix(".stripped.srt")
        strip_srt(pin, pout)
        print(f"[OK] stripped SRT -> {pout}")
        
        # 如果指定了--txt参数，同时导出对白
        if args.txt:
            txt_out = Path(args.txt_out) if args.txt_out else pin.with_suffix(".dialogue.txt")
            count = extract_dialogue_from_srt(pin, txt_out, clean=clean_text)
            print(f"[OK] extracted {count} lines from SRT -> {txt_out}")

    if args.strip_ass:
        pin  = Path(args.strip_ass)
        pout = pin.with_suffix(".stripped.ass")
        strip_ass(pin, pout)
        print(f"[OK] stripped ASS -> {pout}")
        
        # 如果指定了--txt参数，同时导出对白
        if args.txt:
            txt_out = Path(args.txt_out) if args.txt_out else pin.with_suffix(".dialogue.txt")
            count = extract_dialogue_from_ass(pin, txt_out, clean=clean_text)
            print(f"[OK] extracted {count} lines from ASS -> {txt_out}")

    if args.strip_docx:
        pin = Path(args.strip_docx)
        pout = pin.with_suffix(".stripped.docx")
        strip_docx(pin, pout)
        print(f"[OK] stripped DOCX -> {pout}")
        
        # 如果指定了--txt参数，同时导出对白
        if args.txt:
            txt_out = Path(args.txt_out) if args.txt_out else pin.with_suffix(".dialogue.txt")
            count = extract_dialogue_from_docx(pin, txt_out, clean=clean_text)
            print(f"[OK] extracted {count} lines from DOCX -> {txt_out}")

    if args.docx:
        pin = Path(args.docx)
        stats = analyze_docx_yellow(pin, Path(args.csv) if args.csv else None)
        print(json.dumps(stats, ensure_ascii=False, indent=2))
        
        # 如果指定了--txt参数，同时导出全文
        if args.txt:
            txt_out = Path(args.txt_out) if args.txt_out else pin.with_suffix(".dialogue.txt")
            count = extract_dialogue_from_docx(pin, txt_out, clean=clean_text)
            print(f"[OK] extracted {count} lines from DOCX -> {txt_out}")

if __name__ == "__main__":
    main()
